__all__ = ["Chatbot"]

import os
import uuid  # Import uuid to generate unique IDs
from dataclasses import dataclass

import openai
from chromadb import Client
from chromadb.config import Settings
from docx import Document
from ordered_set import OrderedSet
from sentence_transformers import SentenceTransformer

system_prompt = "You are an assistant that analyzes the content of word files containing how-to guides. \
These are general knowledge and onboarding documents. Your job is to answer questions from users and give them \
itemized instructions whenever applicable. If you don't know the answer, just say so."


@dataclass
class Chatbot:
    resources: str
    model: SentenceTransformer = SentenceTransformer("all-MiniLM-L6-v2")  # For embeddings
    client: Client = Client(Settings())

    def __post_init__(self):
        self.collection = self.client.create_collection("knowledge_base")
        self.load_word_files()

    def load_word_files(self):
        for filename in os.listdir(self.resources):
            if filename.endswith(".docx"):
                doc = Document(os.path.join(self.resources, filename))
                self.process_document(doc, filename)

    def process_document(self, doc, filename):
        for para in doc.paragraphs:
            self.learn(para.text, filename)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.learn(cell.text, filename)

    def learn(self, text, filename):
        if text.strip():
            embedding = self.model.encode(text, convert_to_tensor=True).cpu().numpy().tolist()
            unique_id = str(uuid.uuid4())  # Generate a unique ID for each entry

            # Add text and its embedding to the Chroma collection with a unique ID
            self.collection.add(
                documents=[text],
                metadatas=[{"filename": filename}],
                embeddings=[embedding],
                ids=[unique_id],
            )

    def get_response(self, question):
        question_embedding = self.model.encode(question, convert_to_tensor=True).cpu().numpy().tolist()
        relevant_entries, relevant_filenames = self.select_relevant_entries(question_embedding)

        # If no relevant entries are found but relevant filenames are present
        if not relevant_entries and relevant_filenames:
            return (
                "I couldn't find a specific answer, but the following files may contain relevant information:",
                relevant_filenames,
            )

        print(relevant_filenames)

        # If no relevant entries or filenames are found
        if not relevant_entries:
            return "I'm here to help! Please ask a specific question.", []

        # Combine the relevant entries into a context string
        context = "\n".join(relevant_entries)

        # Call the OpenAI API with the context and the user's question
        try:
            response = openai.ChatCompletion.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"{context}\n\nQuestion: {question}"},
                ],
            )
            return response["choices"][0]["message"]["content"], relevant_filenames
        except Exception as e:
            return f"An error occurred: {str(e)}", []

    def select_relevant_entries(self, question_embedding):
        # Query Chroma for relevant entries
        results = self.collection.query(
            query_embeddings=[question_embedding],
            n_results=5,  # Adjust the number of top results as needed
        )

        # Filter based on relevance score
        threshold = 0.8  # Define an appropriate relevance threshold
        relevant_indices = [i for i, distance in enumerate(results["distances"][0]) if distance < threshold]

        if not relevant_indices:
            return [], []

        # Retrieve relevant entries and filenames
        relevant_entries = [results["documents"][0][i] for i in relevant_indices]
        relevant_filenames = list(OrderedSet([results["metadatas"][0][i]["filename"] for i in relevant_indices]))

        return relevant_entries, relevant_filenames
